from docx import Document

doc = Document("Technical_Spesification_Document_(TSD)_-_MPM_Dom_(CR46881).docx")  # ganti dengan nama file DOCX lo
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            if lines:
                print(f"[T{t_idx}R{r_idx}C{c_idx}] {repr(' '.join(lines))}")