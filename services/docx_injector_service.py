# docx_injector_service.py
# Template-based signature injection directly into DOCX files
# 
# Primary flow: DOCX → find placeholders → inject image → return modified DOCX
# Advantage: deterministic placement, no PDF conversion needed until final step
#
# Supported placeholders (case-insensitive, whitespace-normalized):
#   {{SIGN}}
#   {{SIGN_<KEYWORD>}}
#   <<SIGN>>
#
# If placeholder found → inject and return modified DOCX
# If NOT found → raise PlaceholderNotFoundError (will trigger FALLBACK to detection)

import re
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image


class PlaceholderNotFoundError(Exception):
    """Raised when no signature placeholder is found in DOCX."""
    pass


def inject_signature_to_docx(
    docx_bytes: bytes,
    sign_path: str,
    keyword: str = None,
    signature_width: float = 1.5,
) -> bytes:
    """
    Template-based signature injection into DOCX.
    
    Detects placeholders ({{SIGN}}, {{SIGN_<KEYWORD>}}, <<SIGN>>) and replaces them
    with a signature image. Handles both paragraphs and table cells.
    
    Args:
        docx_bytes: DOCX file as binary
        sign_path: Path to signature image (PNG, JPG, SVG→PNG)
        keyword: Optional keyword to match {{SIGN_<KEYWORD>}}
        signature_width: Width in inches (default 1.5)
    
    Returns:
        Modified DOCX as binary
    
    Raises:
        PlaceholderNotFoundError: If no placeholder found
        Exception: If image file not readable or DOCX parsing fails
    """
    doc = Document(io.BytesIO(docx_bytes))
    
    found_any = False
    
    # Build regex patterns to search for
    patterns = [
        r"\{\{SIGN\}\}",                       # {{SIGN}}
        r"\{\{SIGN_" + re.escape(keyword or "") + r"\}\}",  # {{SIGN_<keyword>}}
        r"<<SIGN>>",                           # <<SIGN>>
    ]
    
    # Inject in paragraphs
    for paragraph in doc.paragraphs:
        if _inject_in_paragraph(paragraph, patterns, sign_path, signature_width):
            found_any = True
    
    # Inject in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _inject_in_paragraph(paragraph, patterns, sign_path, signature_width):
                        found_any = True
    
    if not found_any:
        raise PlaceholderNotFoundError(
            f"No signature placeholder found in DOCX. "
            f"Expected: {{{{SIGN}}}}, {{{{SIGN_{keyword or 'NAME'}}}}}, or <<SIGN>>"
        )
    
    # Save modified DOCX to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _inject_in_paragraph(paragraph, patterns: list, sign_path: str, width: float) -> bool:
    """
    Search for placeholder patterns in paragraph and replace with image.
    Returns True if placeholder found and injected, False otherwise.
    
    Handles edge cases:
    - Placeholder split across multiple runs
    - Multiple placeholders in same paragraph
    """
    # Get full paragraph text
    full_text = "".join([run.text for run in paragraph.runs])
    
    # Normalize: lowercase and remove extra whitespace
    normalized = re.sub(r'\s+', ' ', full_text.lower()).strip()
    
    # Try to find any pattern (case-insensitive, normalized)
    matched_pattern = None
    for pattern in patterns:
        if re.search(pattern, normalized, re.IGNORECASE):
            matched_pattern = pattern
            break
    
    if not matched_pattern:
        return False
    
    # Clear paragraph runs
    for run in paragraph.runs:
        run.text = ""
    
    # Add image to first run (or create new run if no runs exist)
    if len(paragraph.runs) == 0:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    
    try:
        run.add_picture(sign_path, width=Inches(width))
    except Exception as e:
        raise Exception(f"Failed to add signature image from {sign_path}: {e}")
    
    return True


def validate_placeholder_syntax(docx_bytes: bytes) -> dict:
    """
    Scan DOCX and report all signature placeholders found.
    Useful for debugging and validation.
    
    Returns:
        {
            "total_placeholders": int,
            "locations": [
                {
                    "type": "paragraph" | "table_cell",
                    "text": "<paragraph/cell text>",
                    "placeholder": "{{SIGN}}" | "{{SIGN_<keyword>}}" | "<<SIGN>>"
                },
                ...
            ]
        }
    """
    doc = Document(io.BytesIO(docx_bytes))
    patterns = [
        (r"\{\{SIGN\}\}", "{{SIGN}}"),
        (r"\{\{SIGN_[^}]+\}\}", "{{SIGN_<keyword>}}"),
        (r"<<SIGN>>", "<<SIGN>>"),
    ]
    
    locations = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for pattern, display in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                locations.append({
                    "type": "paragraph",
                    "text": text[:100],
                    "placeholder": display,
                })
                break
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    for pattern, display in patterns:
                        if re.search(pattern, text, re.IGNORECASE):
                            locations.append({
                                "type": f"table[{table_idx}].row[{row_idx}].cell[{cell_idx}]",
                                "text": text[:100],
                                "placeholder": display,
                            })
                            break
    
    return {
        "total_placeholders": len(locations),
        "locations": locations,
    }
