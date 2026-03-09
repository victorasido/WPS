from docx import Document

doc = Document("Technical_Spesification_Document_(TSD)_-_MPM_Dom_(CR46881).docx")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "Division Head" in cell.text:
                print(f"\n[T{t_idx}R{r_idx}C{c_idx}]")
                for p_idx, para in enumerate(cell.paragraphs):
                    print(f"  para[{p_idx}]: {repr(para.text)}")