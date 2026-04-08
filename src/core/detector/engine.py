"""
services/detector/engine.py
Fungsi utama detect_signature_zones — konduktor yang mengorkestrasi
semua modul detector (docx_xml, matchers) untuk menghasilkan list zona TTD.
"""

from docx import Document
from docx.oxml.ns import qn

from src.shared.text_utils import extract_keyword
from src.infra.config import (
    CONFIDENCE_THRESHOLD,
    LAST_PAGES_SCAN,
    DASH_LINE_MIN,
)
from .matchers import match_cascade, CONF_DASH_BONUS
from .docx_xml import (
    tc_text,
    find_slot_xml,
)


def detect_signature_zones(
    docx_path: str,
    signature_path: str,
    confidence_threshold: float = None,
    last_pages: int = None,
) -> list:
    """
    Scan DOCX dan return list zona TTD berdasarkan keyword dari nama file TTD.

    Keyword = nama file tanpa ekstensi.
    Contoh: "Farino.png"                  → keyword: "Farino"
            "Division Head Divisi IT.png"  → keyword: "Division Head Divisi IT"

    Match cascade (stop di tier pertama yang hit):
        Tier 1 — exact, case-sensitive         → confidence 1.0
        Tier 2 — full phrase, case-insensitive → confidence 0.9
        Tier 3 — partial per kata (≥60% match) → confidence 0.5–0.85

    Slot validation (inject point):
        Priority: blank above same cell → blank above prev row
        Fallback:  blank below same cell → blank below next row
        Skip jika tidak ada blank space sama sekali.
    """
    if confidence_threshold is None:
        confidence_threshold = CONFIDENCE_THRESHOLD
    if last_pages is None:
        last_pages = LAST_PAGES_SCAN

    keyword = extract_keyword(signature_path)
    doc     = Document(docx_path)
    zones   = []

    # ── TABLE SCANNING ──────────────────────────────────────────
    for t_idx, table in enumerate(doc.tables):
        tbl_xml  = table._tbl
        xml_rows = tbl_xml.findall(qn("w:tr"))

        # seen list pakai identity check (tc is x) — id() Python bisa recycle memori
        seen_tcs: list = []

        for r_idx, tr in enumerate(xml_rows):
            xml_cells = tr.findall(qn("w:tc"))

            for c_idx, tc in enumerate(xml_cells):
                if any(tc is s for s in seen_tcs):
                    continue
                seen_tcs.append(tc)

                cell_text = tc_text(tc)

                # Phase 1: match cascade
                match_result = match_cascade(keyword, cell_text)
                if match_result is None:
                    continue

                matched_text, confidence = match_result

                # Phase 2: validate slot
                slot = find_slot_xml(tc, xml_rows, r_idx, c_idx)
                if slot is None:
                    # Tidak ada blank space → skip
                    continue

                inject_p_idx, inject_position, has_dash = slot

                # Bonus confidence jika ada garis ---
                if has_dash:
                    confidence = min(1.0, confidence + CONF_DASH_BONUS)

                # Filter threshold
                if confidence < confidence_threshold:
                    continue

                # Resolve inject row index
                if inject_position == "above_prev_row":
                    inject_r_idx = r_idx - 1
                elif inject_position == "below_next_row":
                    inject_r_idx = r_idx + 1
                else:
                    inject_r_idx = r_idx

                zones.append({
                    "source":          "table",
                    "paragraph_index": 10000 + t_idx * 1000 + r_idx * 10 + c_idx,
                    "table_location":  (t_idx, inject_r_idx, c_idx, inject_p_idx),
                    "matched_name":    matched_text.strip(),
                    "keyword":         keyword,
                    "confidence":      round(confidence, 2),
                    "context":         cell_text.strip()[:80],
                    "inject_position": inject_position,
                })

    # ── PARAGRAPH SCANNING ──────────────────────────────────────
    for p_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue

        match_result = match_cascade(keyword, para_text)
        if match_result is None:
            continue

        matched_text, confidence = match_result

        if confidence < confidence_threshold:
            continue

        zones.append({
            "source":          "paragraph",
            "paragraph_index": p_idx,
            "table_location":  None,
            "matched_name":    matched_text.strip(),
            "keyword":         keyword,
            "confidence":      round(confidence, 2),
            "context":         para_text[:80],
            "inject_position": "below_same",
        })

    zones.sort(key=lambda z: z["paragraph_index"])
    return zones
