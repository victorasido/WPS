from docx import Document
from docx.oxml.ns import qn

doc = Document("Technical_Spesification_Document_(TSD)_-_MPM_Dom_(CR46881).docx")
table = doc.tables[1]
cell = table.rows[4].cells[0]

for i, para in enumerate(cell.paragraphs[:8]):
    pPr = para._p.find(qn("w:pPr"))
    spacing = None
    if pPr is not None:
        spacing = pPr.find(qn("w:spacing"))
    print(f"para[{i}]: '{para.text}' | spacing XML: {spacing.attrib if spacing is not None else 'None'} | style: {para.style.name}")